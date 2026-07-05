from pathlib import Path
from uuid import uuid4

from docx import Document

from app.config import OUTPUT_DIR
from app.state import AgentState


SECTION_HEADINGS = {
    "Executive Summary",
    "Problem Statement",
    "Proposed Solution",
    "Objectives",
    "Scope",
    "Implementation Plan",
    "Timeline",
    "Budget",
    "Budget Assumptions",
    "Expected Benefits",
    "Risks",
    "Risk Assessment",
    "Conclusion",
}


def generate_doc(state: AgentState) -> AgentState:
    """
    Generate a professional Word document
    from the generated business proposal.
    """

    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    document = Document()

    # Main Title
    document.add_heading("Business Proposal", level=1)

    for line in state["proposal"].splitlines():

        line = line.strip()

        if not line:
            continue

        # Markdown Heading
        if line.startswith("#"):
            heading = line.replace("#", "").strip()
            document.add_heading(heading, level=2)
            continue

        # Section Heading
        if line.rstrip(":") in SECTION_HEADINGS:
            document.add_heading(line.rstrip(":"), level=2)
            continue

        # Bullet Points
        if line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(
                line[2:],
                style="List Bullet"
            )
            continue

        # Numbered Lists
        if (
            len(line) > 2
            and line[0].isdigit()
            and line[1] == "."
        ):
            document.add_paragraph(
                line[2:].strip(),
                style="List Number"
            )
            continue

        # Normal Paragraph
        document.add_paragraph(line)

    filename = f"proposal_{uuid4().hex[:8]}.docx"

    filepath = Path(OUTPUT_DIR) / filename

    document.save(filepath)

    state["document"] = str(filepath)

    return state